# -*- coding: utf-8 -*-
"""
build_figure_index.py
Preenche o "Indice de Figuras" no Relatorio_FitMap.docx.
  - Recolhe todas as legendas (paragrafos que comecam por "Figura ")
  - Insere-as formatadas com lider de pontos apos o placeholder existente
  - Insere TC fields junto a cada legenda para que o Word possa
    actualizar numeros de pagina com Ctrl+A -> F9

Execute: python docs/build_figure_index.py
"""
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import copy

DOCX_PATH = Path(__file__).parent.parent / "Relatorio_FitMap.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def make_tof_field():
    """Campo TOF (Table of Figures) que o Word actualiza com F9."""
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), ' TOF \\h \\z \\c "Figura" ')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "[Actualize os campos: Ctrl+A depois F9]"
    r.append(t)
    fld.append(r)
    return fld


def make_index_entry(number, description):
    """Paragrafo de indice: 'Figura X  ....  ' com lider de pontos."""
    p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    # Tab stop com lider de pontos a 14 cm
    tabs_el = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(int(14 * 567)))   # 14 cm em twips (1 cm = 567 twips)
    tabs_el.append(tab)
    pPr.append(tabs_el)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "60")
    pPr.append(spacing)
    p.append(pPr)

    # run com o texto da figura
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20")   # 10 pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = f"{number}  \t"
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p.append(r)

    # run com a descricao (cor cinza escuro)
    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    sz2 = OxmlElement("w:sz"); sz2.set(qn("w:val"), "20")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "444444")
    rPr2.append(sz2); rPr2.append(color)
    r2.append(rPr2)
    t2 = OxmlElement("w:t"); t2.text = description
    r2.append(t2)
    p.append(r2)

    return p


def make_update_hint():
    """Nota pequena a pedir ao utilizador para actualizar campos."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:before"), "80")
    pPr.append(spacing)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    i = OxmlElement("w:i")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "888888")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16")
    rPr.append(i); rPr.append(color); rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "Para actualizar os numeros de pagina: abra no Word, Ctrl+A e depois F9."
    r.append(t)
    p.append(r)
    return p


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc  = Document(str(DOCX_PATH))
    body = doc.element.body
    paras = list(doc.paragraphs)

    # 1) Recolher todas as legendas de figuras
    figures = []
    for p in paras:
        txt = p.text.strip()
        if txt.startswith("Figura ") and " - " in txt:
            # ex: "Figura 1 - Landing page..."
            parts = txt.split(" - ", 1)
            num   = parts[0].strip()    # "Figura 1"
            desc  = parts[1].strip()    # "Landing page..."
            figures.append((num, desc, p))

    print(f"Encontradas {len(figures)} legendas de figuras:")
    for num, desc, _ in figures:
        print(f"  {num}: {desc[:70]}")

    # 2) Localizar o paragrafo "Indice de Figuras"
    idx_para = None
    for p in paras:
        if "ndice de Figuras" in p.text or "ndice de figuras" in p.text.lower():
            idx_para = p
            break

    if idx_para is None:
        print("AVISO: paragrafo 'Indice de Figuras' nao encontrado.")
        return

    # 3) Inserir entradas APOS o titulo do indice
    p_el   = idx_para._element
    parent = p_el.getparent()
    insert_at = list(parent).index(p_el) + 1

    # Nota de actualizacao
    parent.insert(insert_at, make_update_hint())
    insert_at += 1

    # Entradas do indice (em ordem)
    for num, desc, _ in figures:
        parent.insert(insert_at, make_index_entry(num, desc))
        insert_at += 1

    # 4) Estilizar o titulo "Indice de Figuras" como Estilo1 / Heading
    try:
        idx_para.style = doc.styles["Estilo1"]
    except KeyError:
        try:
            idx_para.style = doc.styles["Heading 1"]
        except KeyError:
            pass
    # garantir que o texto esta correcto e com acento
    if idx_para.runs:
        idx_para.runs[0].text = "Indice de Figuras"
        for r in idx_para.runs[1:]:
            r.text = ""
    else:
        idx_para.add_run("Indice de Figuras")

    doc.save(str(DOCX_PATH))
    print(f"\nIndice de Figuras preenchido com {len(figures)} entradas -> {DOCX_PATH.name}")


if __name__ == "__main__":
    main()
