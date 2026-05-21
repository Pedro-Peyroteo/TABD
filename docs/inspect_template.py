"""Inspeciona o template e imprime estrutura completa."""
import sys
from pathlib import Path
from docx import Document

template = Path(__file__).parent.parent / "Template relatório_projeto_IHC.docx"
doc = Document(str(template))

print("=== SECÇÕES ===")
for sec in doc.sections:
    print(f"page={sec.page_width.cm:.1f}x{sec.page_height.cm:.1f}cm | "
          f"margins L={sec.left_margin.cm:.1f} R={sec.right_margin.cm:.1f} "
          f"T={sec.top_margin.cm:.1f} B={sec.bottom_margin.cm:.1f}")

print()
print("=== ESTILOS USADOS ===")
seen = set()
for p in doc.paragraphs:
    if p.style.name not in seen:
        seen.add(p.style.name)
print("  " + "\n  ".join(sorted(seen)))

print()
print(f"=== PARÁGRAFOS ({len(doc.paragraphs)}) ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text[:160].replace("\n", " ")
    if txt.strip() or "Heading" in p.style.name:
        print(f"  [{p.style.name:25}] {txt}")

print()
print(f"=== TABELAS: {len(doc.tables)} ===")
for i, t in enumerate(doc.tables):
    print(f"Tabela {i}: {len(t.rows)} linhas x {len(t.columns)} colunas")
    for r in t.rows[:5]:
        print("  | " + " | ".join(c.text[:40].replace("\n", " ") for c in r.cells))

print()
print("=== FONTE PADRÃO ===")
try:
    style = doc.styles["Normal"]
    f = style.font
    print(f"Nome: {f.name}, tamanho: {f.size}")
except KeyError:
    pass
