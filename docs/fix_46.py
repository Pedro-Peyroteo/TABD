"""
Corrige a seccao 4.6 do Relatorio_FitMap.docx usando body.insert() directo.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOC_PATH = "Relatorio_FitMap.docx"
doc = Document(DOC_PATH)
body = doc.element.body


def get_children():
    return list(body)


def find_idx(fragment):
    for i, el in enumerate(get_children()):
        txt = "".join(t.text or "" for t in el.iter(qn("w:t")))
        if fragment in txt:
            return i
    raise ValueError(f"Nao encontrado: {fragment!r}")


def make_normal(text, italic=False):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "Normal")
    pPr.append(pStyle); p.append(pPr)
    if text:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if italic:
            for tag in ("w:i", "w:iCs"):
                e = OxmlElement(tag); rPr.append(e)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t); p.append(r)
    return p


def make_bullet(text):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "ListParagraph")
    pPr.append(pStyle)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId_el = OxmlElement("w:numId"); numId_el.set(qn("w:val"), "1")
    numPr.append(ilvl); numPr.append(numId_el)
    pPr.append(numPr); p.append(pPr)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t); p.append(r)
    return p


# ── 1. Remover conteudo orphan que esta apos o Heading 1 da seccao 5 ──────────
orphan_fragments = [
    "funcionalidade de geolocaliza",
    "Quando a posi",
    "Move o marcador",
    "Recalcula as dist",
    "arctan2",
    "design divide explicitamente",
    "watch",
    "cancelado",
]

removed = 0
for frag in orphan_fragments:
    try:
        idx = find_idx(frag)
        el = get_children()[idx]
        body.remove(el)
        removed += 1
    except (ValueError, Exception):
        pass
print(f"Removidos {removed} paragrafos orphan")


# ── 2. Inserir conteudo correctamente apos o titulo 4.6 ───────────────────────
# Encontrar posicao do titulo 4.6 no body APOS remocoes
idx_46 = find_idx("4.6. Localiza")
print(f"Titulo 4.6 esta no indice {idx_46}")

# Verificar o que esta imediatamente a seguir
children = get_children()
next_txt = "".join(t.text or "" for t in children[idx_46 + 1].iter(qn("w:t")))
print(f"Proximo elemento: {next_txt[:60].encode('ascii','replace').decode()}")

# Paragrafos a inserir (em ordem directa)
new_elements = [
    make_normal(
        "A funcionalidade de geolocalização GPS foi refinada para suportar tracking "
        "contínuo da posição do utilizador. Em vez da abordagem inicial baseada em "
        "navigator.geolocation.getCurrentPosition() — que captura a posição uma única vez —, "
        "a implementação final recorre a watchPosition(), que regista um listener invocado "
        "automaticamente sempre que o dispositivo reporta uma nova coordenada."
    ),
    make_normal(
        "Quando a posição é actualizada, o sistema executa dois passos sem emitir nova "
        "query à base de dados:"
    ),
    make_bullet(
        "Move o marcador e o círculo de raio no mapa através de setLatLng(), "
        "sem re-renderização completa da camada de marcadores."
    ),
    make_bullet(
        "Recalcula as distâncias de todas as instalações já presentes no painel lateral "
        "usando a fórmula de Haversine implementada em JavaScript:"
    ),
    make_normal(
        "d = 2R · arctan2(√a, √(1−a))     onde     "
        "a = sin²(Δφ/2) + cos φ1 · cos φ2 · sin²(Δλ/2)",
        italic=True
    ),
    make_normal(
        "Este design divide explicitamente as responsabilidades: o MongoDB (via $geoNear) "
        "calcula as distâncias iniciais no momento da query geoespacial; o JavaScript "
        "mantém-nas actualizadas em tempo real sem overhead de rede. A lista de instalações "
        "é reordenada por distância crescente a cada actualização de posição."
    ),
    make_normal(
        "O watch é cancelado (clearWatch) quando o utilizador limpa a seleção, evitando "
        "consumo desnecessário de bateria e eventos de rede."
    ),
]

# Inserir em ordem directa: cada insercao vai para idx_46+1, empurrando as anteriores
# Usamos idx_46 + 1 + offset como posicao de insercao
for offset, el in enumerate(new_elements):
    body.insert(idx_46 + 1 + offset, el)

print(f"Inseridos {len(new_elements)} paragrafos na seccao 4.6")

# ── 3. Verificacao ────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("Ficheiro guardado.")

doc2 = Document(DOC_PATH)
in_46 = False
count = 0
print("\n--- Verificacao seccao 4.6 ---")
for p in doc2.paragraphs:
    t = p.text.strip()
    if "4.6." in t:
        in_46 = True
    if in_46 and t:
        print(f"  [{p.style.name}] {t[:70].encode('ascii','replace').decode()}")
        count += 1
    if "Apresenta" in t and "resultados" in t and in_46:
        break
print(f"Total: {count} paragrafos na seccao 4.6 (incluindo titulo e heading de fecho)")
