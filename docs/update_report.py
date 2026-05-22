"""
Actualiza o Relatorio_FitMap.docx com as melhorias implementadas após a geração inicial:
  - Subsecção 4.6: localização contínua (watchPosition) e Haversine
  - Subsecção 5.1: 234 cidades acessíveis na sidebar com pesquisa
  - Cenário 5.2.4: filtros da sidebar com zoom automático
"""

import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import lxml.etree as etree

DOC_PATH = "Relatorio_FitMap.docx"
doc = Document(DOC_PATH)
paras = doc.paragraphs


# ─── Helpers ──────────────────────────────────────────────────────────────────

def find_para(text_fragment, style=None):
    """Devolve o índice do primeiro parágrafo que contém text_fragment."""
    for i, p in enumerate(paras):
        if text_fragment in p.text:
            if style is None or p.style.name == style:
                return i
    raise ValueError(f"Parágrafo não encontrado: {text_fragment!r}")


def para_xml(p):
    return p._element


def insert_para_after(ref_idx, text, style_name="Normal", bold=False, italic=False):
    """Insere um novo parágrafo após ref_idx. Devolve o novo índice."""
    ref_el = para_xml(paras[ref_idx])
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_name)
    pPr.append(pStyle)
    new_p.append(pPr)

    if text:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            b = OxmlElement("w:b"); rPr.append(b)
        if italic:
            i_el = OxmlElement("w:i"); rPr.append(i_el)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        new_p.append(r)

    ref_el.addnext(new_p)
    # Recarregar lista de parágrafos
    doc.paragraphs  # não invalida — o Document mantém referências ao XML
    return ref_idx + 1


def insert_bullet_after(ref_idx, text):
    """Insere parágrafo bullet (List Paragraph) após ref_idx."""
    ref_el = para_xml(paras[ref_idx])
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "ListParagraph")
    pPr.append(pStyle)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "1")
    numPr.append(ilvl); numPr.append(numId)
    pPr.append(numPr)
    new_p.append(pPr)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    new_p.append(r)
    ref_el.addnext(new_p)
    return ref_idx + 1


# ─── 1. Subsecção 4.6 após o fim da 4.5 ──────────────────────────────────────
# Localizar o último parágrafo da secção 4.5 (antes do Heading 1 da secção 5)

idx_sec5 = find_para("Apresentação e análise dos resultados", style="Heading 1")

# Inserir de baixo para cima (cada insert desloca os índices seguintes)
# Vamos construir a lista de parágrafos a inserir e inserir em ordem

new_paras_46 = [
    ("Estilo2", "4.6. Localização contínua e recálculo de distâncias no cliente", False),
    ("Normal",
     "A funcionalidade de geolocalização GPS foi refinada para suportar tracking "
     "contínuo da posição do utilizador. Em vez da abordagem inicial baseada em "
     "navigator.geolocation.getCurrentPosition() — que captura a posição uma única vez —, "
     "a implementação final recorre a watchPosition(), que regista um listener que é "
     "invocado automaticamente sempre que o dispositivo reporta uma nova coordenada.", False),
    ("Normal",
     "Quando a posição é actualizada, o sistema executa dois passos sem emitir nova "
     "query à base de dados:", False),
]

bullet_46 = [
    "Move o marcador e o círculo de raio no mapa através de setLatLng(), sem re-renderização completa.",
    "Recalcula as distâncias de todas as instalações já presentes no painel lateral "
    "usando a fórmula de Haversine implementada em JavaScript no cliente:",
]

formula_text = (
    "d = 2R · arctan2(√a, √(1−a))    onde    "
    "a = sin²(Δφ/2) + cos φ₁ · cos φ₂ · sin²(Δλ/2)"
)

close_46 = [
    ("Normal",
     "Este design divide explicitamente as responsabilidades: o MongoDB (via $geoNear) "
     "calcula as distâncias iniciais no momento da query geoespacial; o JavaScript "
     "mantém-nas actualizadas em tempo real sem overhead de rede. A lista de instalações "
     "é reordenada por distância crescente a cada actualização de posição.", False),
    ("Normal",
     "O watch é cancelado (clearWatch) quando o utilizador limpa a selecção, evitando "
     "consumo desnecessário de bateria e eventos de rede.", False),
]

# Inserir em sequência (cada insert empurra os índices, mas como inserimos sempre
# "após o mesmo ref_idx" o resultado é ordem inversa). Para obter ordem directa,
# usamos um cursor que avança.

cursor = idx_sec5 - 1  # último parágrafo antes da secção 5

for style, text, bold in new_paras_46:
    insert_para_after(cursor, text, style_name=style, bold=bold)
    cursor += 1

for bullet in bullet_46:
    insert_bullet_after(cursor, bullet)
    cursor += 1

insert_para_after(cursor, formula_text, style_name="Normal", italic=True)
cursor += 1

for style, text, bold in close_46:
    insert_para_after(cursor, text, style_name=style, bold=bold)
    cursor += 1


# ─── 2. Actualizar secção 5.1 ─────────────────────────────────────────────────
# Encontrar parágrafo sobre 234 cidades e adicionar frase sobre a sidebar

idx_234 = find_para("Foram registadas 234 cidades distintas")
para_234 = paras[idx_234]
# Append frase ao texto existente
run = para_234.add_run(
    " A totalidade destas 234 cidades está acessível na barra lateral da aplicação, "
    "com scroll e pesquisa incremental, permitindo ao utilizador filtrar instalações "
    "por qualquer município sem necessidade de efectuar pesquisa no mapa."
)


# ─── 3. Novo cenário 5.2.4 ────────────────────────────────────────────────────
# Inserir após o fim do cenário 5.2.3 (antes de 5.3)

idx_53 = find_para("5.3. Análise de performance")

new_paras_524 = [
    ("Normal", "5.2.4.  Cenário: «filtrar ginásios e ver apenas os marcadores relevantes»", False),
    ("Normal", "", False),  # linha em branco
]

bullets_524 = [
    "O utilizador seleciona o pill «Ginásio» na secção CATEGORIA da barra lateral. "
    "O cliente emite GET /api/facilities?category=Ginásio&limit=3000 ao backend.",
    "O servidor aplica {$match: {category: \"Ginásio\"}} antes do $project e do $limit, "
    "devolvendo apenas os 708 ginásios registados.",
    "O mapa remove todos os marcadores anteriores e redesenha apenas os 708 marcadores "
    "de ginásio, com cor característica (#ff5b3a). O viewport faz fitBounds automático "
    "sobre o conjunto filtrado, tornando a filtragem visualmente evidente.",
    "Um toast de confirmação apresenta «708 instalações · Ginásio». O contador VIS. "
    "no cabeçalho da aplicação actualiza em tempo real.",
    "O utilizador pode combinar o filtro de categoria com um filtro de modalidade "
    "(e.g., «fitness») e/ou uma cidade da sidebar, sendo cada combinação resolvida "
    "por uma única query MongoDB com múltiplos predicados no $match.",
]

close_524 = [
    ("Normal", "", False),
]

cursor2 = idx_53 - 1

for style, text, bold in new_paras_524:
    insert_para_after(cursor2, text, style_name=style, bold=bold)
    cursor2 += 1

for bullet in bullets_524:
    insert_bullet_after(cursor2, bullet)
    cursor2 += 1

for style, text, bold in close_524:
    insert_para_after(cursor2, text, style_name=style, bold=bold)
    cursor2 += 1


# ─── Guardar ──────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
print("Relatório actualizado com sucesso.")
print(f"  - Subsecção 4.6 adicionada")
print(f"  - Secção 5.1 actualizada (234 cidades na sidebar)")
print(f"  - Cenário 5.2.4 adicionado (filtros com zoom automático)")
